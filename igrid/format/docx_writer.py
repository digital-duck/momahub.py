"""Convert markdown text to a .docx file using python-docx."""
from __future__ import annotations

import re
from pathlib import Path


def markdown_to_docx(md_text: str, output_path: str, title: str = "") -> Path:
    """Write *md_text* (simple markdown) to a Word document at *output_path*."""
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            _add_runs(p, stripped)

    out = Path(output_path)
    doc.save(str(out))
    return out


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CODE_RE = re.compile(r"`([^`]+)`")


def _add_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, handling **bold** and `code` spans."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`", text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(1) is not None:
            paragraph.add_run(m.group(1)).bold = True
        else:
            run = paragraph.add_run(m.group(2))
            run.font.name = "Courier New"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
